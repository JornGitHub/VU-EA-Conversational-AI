"""Field catalog support for Inschrijvingen_aggr_UNL_2025.csv.

The catalog is built from the primary Word document and then used as a hard
source for field-level questions.  Keep this module dependency-light except for
python-docx, which is already required by the ingestion pipeline.
"""
from __future__ import annotations

import json
import re
from pathlib import Path
from typing import Any


PROJECT_ROOT = Path(__file__).resolve().parents[2]
DATA_DIR = PROJECT_ROOT / "data"
PRIMARY_SOURCE_DOCUMENT = "Aggregaatbestand inschrijvingen_1cHO2025.docx"
PRIMARY_DATASET = "Inschrijvingen_aggr_UNL_2025.csv"
PRIMARY_SOURCE_PATH_HINTS = [
    "sources/1cHO Documentatie/Aggregaatbestand inschrijvingen_1cHO2025.docx",
    "1cHO Documentatie/Aggregaatbestand inschrijvingen_1cHO2025.docx",
]
FIELD_CATALOG_PATH = DATA_DIR / "inschrijvingen_aggr_2025_field_catalog.json"
PSEUDO_GOLD_PATH = DATA_DIR / "gold_standard_inschrijvingen_aggr_2025.jsonl"

FIELD_NAME_OVERRIDES = {
    33: "Eerstejaars aan deze instelling",
    48: "Indicatie EER op peildatum 1 oktober",
    52: "Nationaliteit 1 op peildatum 1 oktober",
    53: "Indicatie internationale student op peildatum 1 oktober",
    54: "Aantal",
}

GLOBAL_TRANSFORMATIONS = [
    "Eerste-jaarvelden krijgen 0001 als ze gelijk zijn aan Inschrijvingsjaar, anders 0000.",
    "Diverse eerstejaarsindicaties zetten waarden 2, 4 en 5 om naar 6.",
    "Diverse soort-inschrijvingsvelden zetten waarden ongelijk aan 1, 2 en 4 om naar 6.",
    "Records zijn daarna geaggregeerd op alle geselecteerde velden.",
]
SELECTION_INFO = {
    "based_on": "1cHO2025",
    "records": [
        "Actuele instelling = 21PB, 21PC, 21PD, 21PE, 21PF, 21PG, 21PH, 21PI, 21PJ, 21PK, 21PL, 21PM, 21PN",
        "Soort inschrijving soort ho = 1, 2, 3, 4",
    ],
    "limitations": [
        "CBS levert alleen de laatste 15 jaar.",
        "Alleen ROD wordt gebruikt; daardoor ontbreekt Indicatie soort programma en is Liberal Arts and Sciences minder goed uit te splitsen.",
    ],
    "changes_vs_previous_delivery": [
        "Ten opzichte van 1cHO2024 zijn Nationaliteit 1 op peildatum 1 oktober en Indicatie internationale student op peildatum 1 oktober toegevoegd.",
    ],
}

ALIASES = {
    "Inschrijvingsjaar": ["jaar"],
    "Maand vanaf": ["startmaand"],
    "Indicatie actief op peildatum": ["actief op peildatum", "peildatum 1 oktober"],
    "Indicatie EER actueel": ["eer", "EER", "EU/EER", "EER-student", "eer actueel"],
    "Indicatie EER op peildatum 1 oktober": ["eer peildatum", "EER op peildatum", "EU/EER peildatum"],
    "Opleiding actueel equivalent": ["actuele opleiding", "opleiding actueel", "actueel equivalent", "opleiding code actueel"],
    "Opleiding historisch equivalent": ["historische opleiding", "opleiding historisch", "historisch equivalent", "opleiding code historisch"],
    "Indicatie internationale student": ["internationale student", "international student"],
    "Indicatie internationale student op peildatum 1 oktober": ["internationale student peildatum", "international student peildatum"],
    "Nationaliteit 1": ["nationaliteit actueel"],
    "Nationaliteit 1 op peildatum 1 oktober": ["nationaliteit peildatum"],
    "Aantal": ["telveld", "count"],
}
COMMON_ALIASES = ["veld", "variabele", "kolom", "field"]


def normalize_text(text: Any) -> str:
    return re.sub(r"\s+", " ", re.sub(r"[^\w]+", " ", str(text).lower())).strip()


def find_primary_source(root: Path = PROJECT_ROOT) -> Path | None:
    for hint in PRIMARY_SOURCE_PATH_HINTS:
        p = root / hint
        if p.exists():
            return p
    matches = list(root.glob(f"**/{PRIMARY_SOURCE_DOCUMENT}"))
    return matches[0] if matches else None


def clean_field_name(name: str, number: int) -> str:
    if number in FIELD_NAME_OVERRIDES:
        return FIELD_NAME_OVERRIDES[number]
    name = re.sub(r"\s+", " ", name).strip()
    name = re.sub(r"\s*\(NIEUW\)", "", name, flags=re.I)
    return name.strip()


def parse_possible_values(lines: list[str]) -> tuple[list[dict[str, str]], list[str], list[str]]:
    values: list[dict[str, str]] = []
    notes: list[str] = []
    refs: list[str] = []
    for line in lines:
        s = line.strip()
        if not s:
            continue
        if s.lower().startswith(("nb", "(*)")):
            notes.append(s)
        # Keep concrete referenced helper filenames, not prose lines or old dataset examples.
        m = re.match(r"^(\[leeg\]|[\w/> ]+?)\s*=\s*(.+)$", s)
        if m:
            values.append({"code": m.group(1).strip(), "meaning": m.group(2).strip()})
        elif re.match(r"^\d{2}\s+t/m\s+\d{2}$", s, re.I):
            values.append({"code": s, "meaning": "toegestane maandwaarden"})
        elif s.startswith(">"):
            values.append({"code": s.split()[0], "meaning": " ".join(s.split()[1:])})
    return values, notes, refs


def build_catalog(source_path: Path | None = None) -> list[dict[str, Any]]:
    source_path = source_path or find_primary_source()
    if source_path is None:
        raise FileNotFoundError(PRIMARY_SOURCE_DOCUMENT)
    # python-docx kost ± 200 ms aan importtijd en is alleen hier nodig, bij het
    # bouwen van de catalogus - niet bij het beantwoorden van een vraag.
    from docx import Document

    doc = Document(source_path)
    table = doc.tables[0]
    rows = []
    for r in table.rows[1:]:
        cells = [c.text.strip() for c in r.cells]
        if cells and cells[0].isdigit():
            n = int(cells[0])
            rows.append((n, clean_field_name(cells[1], n), cells[2], cells[3]))

    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    start = paragraphs.index("4. Veldbeschrijving") + 1
    heading_by_norm = {normalize_text(name): (n, name) for n, name, *_ in rows}
    aliases_to_heading = {"indicatie eer peildatum": "Indicatie EER op peildatum 1 oktober", "verblijfsjaar actuele opleiding instelling": "Verblijfsjaar Actuele Opleiding-Instelling"}
    blocks: dict[int, list[str]] = {}
    current: int | None = None
    for line in paragraphs[start:]:
        cleaned = re.sub(r"\s*\(Vanaf.*?\)$", "", line)
        cleaned = re.sub(r"\s*\(NIEUW\)$", "", cleaned)
        norm = normalize_text(cleaned)
        target = heading_by_norm.get(norm)
        if not target and norm in aliases_to_heading:
            target = heading_by_norm.get(normalize_text(aliases_to_heading[norm]))
        if target:
            current = target[0]
            blocks[current] = []
        elif current is not None:
            blocks[current].append(line)

    catalog = []
    for n, name, bron, type_field in rows:
        block = blocks.get(n, [])
        possible_idx = next((i for i, line in enumerate(block) if line.lower().startswith("mogelijke waarden")), len(block))
        desc_lines = [line for line in block[:possible_idx] if line.strip(" .")]
        rest = block[possible_idx + 1 :] if possible_idx < len(block) else []
        possible, notes, refs = parse_possible_values(rest)
        refs.extend(ref for ref in re.findall(r"\b[\w-]+\.(?:csv|txt|asc)\b", " ".join(block), flags=re.I) if not ref.lower().startswith("inschrijvingen_aggr_unl_"))
        transforms = []
        low_name = normalize_text(name)
        if low_name.startswith("soort inschrijving"):
            transforms.append("Waarden ongelijk aan 1, 2 en 4 zijn vóór aggregatie omgezet naar 6.")
        if "eerstejaar" in low_name or low_name.startswith("eerste jaar"):
            transforms.append("Eerstejaars-/eerste-jaarvelden zijn vóór aggregatie geharmoniseerd; eerste-jaarvelden zijn 0001 bij gelijkheid aan Inschrijvingsjaar en anders 0000, indicaties 2/4/5 zijn omgezet naar 6.")
        if name == "Geslacht":
            transforms.append("1cHO-waarden M en V zijn omgezet naar 1 en 2.")
        related = [other for _, other, *_ in rows if other != name and any(tok in normalize_text(other) for tok in normalize_text(name).split()[:2])][:8]
        if name == "Opleiding actueel equivalent":
            related = ["Opleiding historisch equivalent", "Iscedf2013rubriek", "Croho-onderdeel actuele opleiding"] + [r for r in related if r not in {"Opleiding historisch equivalent", "Iscedf2013rubriek", "Croho-onderdeel actuele opleiding"}]
        if name == "Opleiding historisch equivalent":
            related = ["Opleiding actueel equivalent", "Iscedf2013rubriek", "Croho-onderdeel actuele opleiding"] + [r for r in related if r not in {"Opleiding actueel equivalent", "Iscedf2013rubriek", "Croho-onderdeel actuele opleiding"}]
        if name == "Indicatie internationale student":
            desc_lines.append("Definitie internationale student: geen Nederlandse nationaliteit en geen Nederlandse vooropleiding voor het HO. De actuele variant gebruikt de actuele eerste nationaliteit en kan door naturalisatie met terugwerkende kracht wijzigen.")
        if name == "Indicatie internationale student op peildatum 1 oktober":
            desc_lines.append("Definitie internationale student: geen Nederlandse nationaliteit en geen Nederlandse vooropleiding voor het HO. De peildatumvariant gebruikt de eerste nationaliteit op peildatum 1 oktober; jaren vóór naturalisatie blijven behouden als de student toen internationale student was.")
        aliases = COMMON_ALIASES + ALIASES.get(name, [])
        entry = {
            "field_number": n,
            "field_name": name,
            "normalized_field_name": normalize_text(name),
            "dataset": PRIMARY_DATASET,
            "source_document": PRIMARY_SOURCE_DOCUMENT,
            "source_path": source_path.relative_to(PROJECT_ROOT).as_posix() if source_path.is_relative_to(PROJECT_ROOT) else source_path.as_posix(),
            "bron": bron,
            "type_field": type_field,
            "description": " ".join(desc_lines).strip() or "Geen afzonderlijke beschrijving gevonden in de veldbeschrijving.",
            "possible_values": possible,
            "notes": notes,
            "references": sorted(set(refs), key=str.lower),
            "transformations": transforms,
            "related_fields": related,
            "aliases": aliases,
            "gold_questions": [
                f"Wat betekent veld {name}?",
                f"Wat is de definitie van {name}?",
                f"Welke mogelijke waarden heeft {name}?",
                f"In welk bestand staat {name}?",
                f"Wat is de bron en type veld van {name}?",
                f"Wat zijn de aandachtspunten/NB's bij {name}?",
                f"Toon alles over veld {name}.",
            ],
            "source_metadata": {"document_part": "4. Veldbeschrijving", "layout_table_row": n},
        }
        catalog.append(entry)
    return catalog


def write_catalog_and_gold(dry_run: bool = False) -> dict[str, Any]:
    catalog = build_catalog()
    cases = []
    for entry in catalog:
        must = [entry["field_name"], PRIMARY_DATASET]
        if entry["possible_values"]:
            must += [str(entry["possible_values"][0]["code"])]
        if "peildatum" in entry["normalized_field_name"] or "internationale student" in entry["normalized_field_name"]:
            must.append("peildatum 1 oktober")
        if "internationale student" in entry["normalized_field_name"]:
            must += ["geen Nederlandse nationaliteit", "geen Nederlandse vooropleiding", "voor het HO"]
        if "internationale student op peildatum" in entry["normalized_field_name"]:
            must += ["J", "N", "naturalisatie"]
        for q in entry["gold_questions"]:
            cases.append({
                "query": q,
                "expected_field": entry["field_name"],
                "expected_dataset": PRIMARY_DATASET,
                "expected_source_document": PRIMARY_SOURCE_DOCUMENT,
                "expected_intent": "possible_values" if "mogelijke waarden" in q.lower() else "field_detail",
                "must_include": list(dict.fromkeys(must)),
                "should_not_confuse_with": [f["field_name"] for f in catalog if f["field_name"] != entry["field_name"] and normalize_text(entry["field_name"]) in normalize_text(f["field_name"])][:2],
            })
    if not dry_run:
        DATA_DIR.mkdir(exist_ok=True)
        FIELD_CATALOG_PATH.write_text(json.dumps(catalog, ensure_ascii=False, indent=2), encoding="utf-8")
        PSEUDO_GOLD_PATH.write_text("\n".join(json.dumps(c, ensure_ascii=False) for c in cases) + "\n", encoding="utf-8")
    return {
        "fields_found": len(catalog),
        "field_descriptions": sum(1 for e in catalog if e["description"]),
        "possible_values": sum(len(e["possible_values"]) for e in catalog),
        "notes": sum(len(e["notes"]) for e in catalog),
        "pseudo_gold_cases": len(cases),
        "complete_54_fields": len(catalog) == 54,
    }


def load_catalog(path: Path = FIELD_CATALOG_PATH) -> list[dict[str, Any]]:
    """Return the inschrijvingen field catalog, cached until the file changes."""
    if not path.exists():
        return []
    from src.definitions.corpus import cached_json

    return cached_json(path)
