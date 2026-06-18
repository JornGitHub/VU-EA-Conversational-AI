import re, json, unicodedata
from pathlib import Path
from collections import defaultdict
from docx import Document
import fitz

PROJECT_ROOT = Path(__file__).resolve().parent
BASE = PROJECT_ROOT / 'sources' / '1cHO Documentatie'
OUT = PROJECT_ROOT / 'data'
OUT.mkdir(parents=True, exist_ok=True)

def clean(s):
    s = s.replace('\xa0', ' ')
    s = re.sub(r'[ \t]+', ' ', s)
    s = re.sub(r'\n{3,}', '\n\n', s)
    return s.strip()

def slug(s):
    s = s.lower()
    s = unicodedata.normalize('NFKD', s).encode('ascii','ignore').decode('ascii')
    s = re.sub(r'[^a-z0-9]+','_',s).strip('_')
    return s[:90] or 'entry'

def parse_values(text):
    vals=[]
    in_vals=False
    for line in text.splitlines():
        l=line.strip()
        if re.match(r'^Mogelijke waarden:?$', l, re.I):
            in_vals=True; continue
        if in_vals:
            if not l: continue
            m=re.match(r'^(.+?)\s*=\s*(.+)$', l)
            if m:
                vals.append({'code': clean(m.group(1)), 'label': clean(m.group(2))})
            else:
                # stop once no value-like pattern after a few codes
                if vals and (len(l)>80 or re.match(r'^(NB|Zie|De|Dit|Wanneer|Het|Deze|Als)\b',l)):
                    break
    return vals

def infer_tags(term, definition=''):
    text=(term+' '+definition).lower()
    tags=[]
    for key, tag in [
        ('internation', 'internationalisering'), ('eer', 'eer/nationaliteit'), ('nationaliteit','nationaliteit'),
        ('inschrijv', 'inschrijving'), ('instroom','instroom'), ('diploma','diploma'), ('examen','diploma'),
        ('cohort','cohort'), ('eoi','cohort'), ('her', 'herinschrijving'), ('studiesucces','studievoortgang'),
        ('uitval','studievoortgang'), ('studiewissel','studievoortgang'), ('studieduur','studievoortgang'),
        ('vooropleiding','vooropleiding'), ('opleiding','opleiding'), ('instelling','instelling'), ('croho','opleiding'),
        ('geslacht','studentkenmerk'), ('leeftijd','studentkenmerk'), ('geboorteland','studentkenmerk'),
        ('generatie','studentkenmerk'), ('herkomst','studentkenmerk'), ('verblijfsjaar','verblijfsjaar')]:
        if key in text and tag not in tags: tags.append(tag)
    return tags

def source_dataset_from_docx(path):
    stem=path.name
    mapping={
        "Aggregaatbestand inschrijvingen": "Inschrijvingen_aggr_UNL_2025.csv",
        "Aggregaatbestand diploma": "Diplomas_aggr_UNL_2025.csv",
        "EOIcohort_aggr": "EOIcohort_aggr_UNL_2025.csv",
        "EOIcohort_UNL": "EOIcohort_UNL_2025.csv / EOIcohort_21P*_2025.csv",
        "Examencohort (aggr)": "Gediplomeerdencohort_aggr_UNL_2025.csv",
        "Examencohort_UNL": "Gediplomeerdencohort_UNL_2025.csv / Gediplomeerdencohort_21P*_2025.csv",
    }
    for k,v in mapping.items():
        if k in stem: return v
    return None

def extract_docx(path):
    entries=[]
    doc=Document(path)
    paragraphs=[clean(p.text) for p in doc.paragraphs if clean(p.text)]
    # table field metadata
    table_fields=[]
    if doc.tables:
        for table in doc.tables:
            rows=[]
            for row in table.rows:
                rows.append([clean(cell.text.replace('\n',' / ')) for cell in row.cells])
            header=' | '.join(rows[0]).lower() if rows else ''
            # Find column name for Naam veld, bron, type
            if rows:
                header_row=rows[0]
                name_idx=None; bron_idx=None; type_idx=None; nr_idx=None
                # sometimes first two rows combined; locate in first two rows
                hdr_combined=[]
                for c in range(len(rows[0])):
                    joined=' '.join(rows[r][c] for r in range(min(2,len(rows))))
                    hdr_combined.append(joined.lower())
                for i,h in enumerate(hdr_combined):
                    if 'naam veld' in h or h.strip()=='naam veld': name_idx=i
                    if h.strip()=='bron' or ' bron' in h: bron_idx=i
                    if 'type veld' in h: type_idx=i
                    if 'variabel' in h: nr_idx=i
                if name_idx is None:
                    continue
                for r in rows[1:]:
                    if name_idx < len(r):
                        name=clean(r[name_idx])
                        if not name or name.lower() in ['naam veld'] or 'naam veld' in name.lower():
                            continue
                        table_fields.append({
                            'name': name,
                            'source': r[bron_idx] if bron_idx is not None and bron_idx<len(r) else '',
                            'type': r[type_idx] if type_idx is not None and type_idx<len(r) else '',
                            'variable_number': r[nr_idx] if nr_idx is not None and nr_idx<len(r) else ''
                        })
    dataset=source_dataset_from_docx(path)
    for f in table_fields:
        entries.append({
            'term': f['name'],
            'definition': '',
            'entry_type':'field_index',
            'source_file': path.name,
            'source_type':'docx',
            'dataset_or_file': dataset,
            'field_name': f['name'],
            'field_source': f.get('source',''),
            'field_type': f.get('type',''),
            'variable_number': f.get('variable_number',''),
            'values': [],
            'aliases': [],
            'tags': infer_tags(f['name']),
            'confidence':'high',
        })
    # definition parsing from 4. Veldbeschrijving onward
    start=0
    for i,p in enumerate(paragraphs):
        if re.match(r'^4\.\s*Veldbeschrijving', p, re.I):
            start=i+1; break
    paras=paragraphs[start:]
    def is_heading(p):
        if len(p)>140 or len(p)<2: return False
        if not (p[0].isupper() or p[0].isdigit() or p[0] in "\"'("): return False
        if p.endswith('.') or p.endswith(':') or p.endswith(';'): return False
        if p.startswith('-') or p.startswith('•'): return False
        if '=' in p: return False
        if re.match(r'^(Mogelijke waarden|NB!?|Zie|Deze|Dit|De |Het |Een |Voor |Wanneer|Als |Bij |Omdat|Daarnaast|Records|Velden|Ten opzichte|Twee variabelen|Pagina|Vereniging)', p, re.I): return False
        # Skip obvious table values / sentence fragments
        if re.match(r'^\d+\s*=|^[A-Z0-9\[\]/]+\s*=', p): return False
        # likely if close to known field, or just field-like noun phrase
        return True
    # collect blocks
    current=None; body=[]
    for p in paras:
        if is_heading(p):
            if current and body:
                definition=clean('\n'.join(body))
                entries.append({
                    'term': current,
                    'definition': definition,
                    'entry_type':'field_definition',
                    'source_file': path.name,
                    'source_type':'docx',
                    'dataset_or_file': dataset,
                    'field_name': current,
                    'values': parse_values(definition),
                    'aliases': [],
                    'tags': infer_tags(current, definition),
                    'confidence':'medium',
                })
            current=p; body=[]
        else:
            if current:
                body.append(p)
    if current and body:
        definition=clean('\n'.join(body))
        entries.append({
            'term': current,
            'definition': definition,
            'entry_type':'field_definition',
            'source_file': path.name,
            'source_type':'docx',
            'dataset_or_file': dataset,
            'field_name': current,
            'values': parse_values(definition),
            'aliases': [],
            'tags': infer_tags(current, definition),
            'confidence':'medium',
        })
    return entries

def parse_txt_1cijfer(path):
    entries=[]
    text=path.read_text(encoding='cp1252')
    dataset='1cyferho_2025_v1.0.asc'
    # layout fields
    m=re.search(r'Lay-out bestand\n=+\n(.*?)\n\(\*\)', text, flags=re.S)
    if m:
        for line in m.group(1).splitlines():
            mm=re.match(r'^(.+?)\s{2,}(\d+)\s+(\d+)\s*$', line.rstrip())
            if mm:
                entries.append({
                    'term': clean(mm.group(1)), 'definition':'', 'entry_type':'field_layout',
                    'source_file':path.name,'source_type':'txt','dataset_or_file':dataset,
                    'field_name':clean(mm.group(1)), 'start_position': int(mm.group(2)), 'length': int(mm.group(3)),
                    'values':[], 'aliases':[], 'tags':infer_tags(mm.group(1)), 'confidence':'high'})
    # field descriptions using underlined headings
    vm=re.search(r'Veldbeschrijving\n=+\n(.*)$', text, flags=re.S)
    if vm:
        sec=vm.group(1)
        matches=list(re.finditer(r'(?m)^(.+?)\n-{3,}\n', sec))
        for i,match in enumerate(matches):
            term=clean(match.group(1))
            body_start=match.end()
            body_end=matches[i+1].start() if i+1<len(matches) else len(sec)
            body=clean(sec[body_start:body_end])
            if body:
                entries.append({
                    'term':term,'definition':body,'entry_type':'field_definition',
                    'source_file':path.name,'source_type':'txt','dataset_or_file':dataset,
                    'field_name':term,'values':parse_values(body),'aliases':[],
                    'tags':infer_tags(term,body),'confidence':'high'})
    return entries

def parse_txt_vakken(path):
    entries=[]
    text=path.read_text(encoding='cp1252')
    # Add intro as concept
    intro=text.split('Decodeertabel vakcode')[0]
    entries.append({'term':'Vakkenbestanden','definition':clean(intro),'entry_type':'concept_definition',
                    'source_file':path.name,'source_type':'txt','dataset_or_file':'Dec_vakcode.asc / vakgegevens',
                    'field_name':'','values':[],'aliases':['vakgegevens','cijferlijst','vakkenbestand'],
                    'tags':['vooropleiding'],'confidence':'high'})
    current_section=None
    for line in text.splitlines():
        mm=re.match(r'^(.+?)\s{2,}(\d+)\s+(\d+)(?:\s+(.*))?$', line.rstrip())
        if mm and not any(x in mm.group(1).lower() for x in ['startpositie','aantal posities']):
            term=clean(mm.group(1))
            opm=clean(mm.group(4) or '')
            entries.append({'term':term,'definition':opm,'entry_type':'field_layout',
                    'source_file':path.name,'source_type':'txt','dataset_or_file':'Dec_vakcode.asc / vakgegevens',
                    'field_name':term,'start_position': int(mm.group(2)), 'length': int(mm.group(3)),
                    'values':[],'aliases':[],'tags':infer_tags(term,opm),'confidence':'high'})
    return entries

def parse_handleiding_pdf(path):
    entries=[]
    doc=fitz.open(path)
    text='\n'.join(doc[i].get_text() for i in range(4,16)) # pages 5-16 definitions/variables
    # remove repeated headers/footers
    lines=[]
    for l in text.splitlines():
        l=clean(l)
        if not l: continue
        if l in ['Handleiding studentenaantallen en studievoortgang','Vereniging Hogescholen'] or l.startswith('Versie ') or l.startswith('Pagina '): continue
        if re.match(r'^\d+ https?://',l): continue
        lines.append(l)
    def is_heading(l):
        if len(l)>100: return False
        if not (l[0].isupper() or l[0].isdigit() or l[0] in "\"'("): return False
        if l.endswith('.') or l.endswith(':'): return False
        if l.startswith('-') or l.startswith('o') or l.startswith(''): return False
        if re.match(r'^(Voorbeeld|Hieronder|Meer informatie|Alleen|I\.v\.m\.|Kengetallen|In aanvulling|Soms wordt|Dit betekent|Afhankelijk|Een student|Bij |Per |Standaard |Wanneer |Velden|De variabele|Het bevat|Deze variabele|Informatie|Studenten|Voor |\d+)',l): return False
        if '=' in l: return False
        # page section/term headings; include parentheses as variables
        return True
    cur=None; body=[]
    for l in lines:
        if is_heading(l):
            if cur and body:
                definition=clean('\n'.join(body))
                if len(definition)>20:
                    entries.append({'term':cur,'definition':definition,'entry_type':'concept_definition',
                        'source_file':path.name,'source_type':'pdf','dataset_or_file':'VH informatieproducten / 1cijferHO',
                        'field_name':'','values':parse_values(definition),'aliases':[],
                        'tags':infer_tags(cur,definition),'confidence':'medium'})
            cur=l; body=[]
        else:
            if cur: body.append(l)
    if cur and body:
        definition=clean('\n'.join(body))
        if len(definition)>20:
            entries.append({'term':cur,'definition':definition,'entry_type':'concept_definition',
                        'source_file':path.name,'source_type':'pdf','dataset_or_file':'VH informatieproducten / 1cijferHO',
                        'field_name':'','values':parse_values(definition),'aliases':[],
                        'tags':infer_tags(cur,definition),'confidence':'medium'})
    return entries

def parse_trendrapport_pdf(path):
    entries=[]
    doc=fitz.open(path)
    # Appendix starts on page 143 (1-based); collect lines
    lines=[]
    for i in range(142, doc.page_count):
        for l in doc[i].get_text().splitlines():
            l=clean(l)
            if not l: continue
            if l.startswith('Trendrapport HO') or l.startswith('Copyright') or l == 'Appendix: definities en tabellen': continue
            if re.match(r'^\d+$', l): continue
            lines.append(l)
    # Indicator title lines start 1.1.1 / 6.2.10 etc.
    title_re=re.compile(r'^(\d+(?:\.\d+){1,3})\s+(.{8,})$')
    year_re=re.compile(r'^(\d{2}/\d{2}|\d{4}|[Qq]\d|[0-9]+[,\.]?[0-9]*%|\[TR|TR\d|hbo|wo|totaal)')
    i=0
    while i < len(lines):
        m=title_re.match(lines[i])
        if not m:
            i+=1; continue
        number,title=m.group(1),m.group(2)
        body=[]; j=i+1
        while j < len(lines):
            if title_re.match(lines[j]):
                break
            # stop at table code, years, obvious table rows after having body
            if body and (lines[j].startswith('[TR') or re.match(r'^\d{2}/\d{2}$',lines[j]) or re.match(r'^[\d\.]+%$', lines[j]) or lines[j].lower() in ['hbo','wo','totaal']):
                # skip table until next title
                while j < len(lines) and not title_re.match(lines[j]):
                    j+=1
                break
            # stop if line is table-like and body sufficiently started
            if body and (re.match(r'^[\d\.]+$', lines[j]) or re.match(r'^\d+[,.]\d+%$', lines[j])):
                j+=1; continue
            # collect definition text; ignore very short table fragments when no sentence
            if not lines[j].startswith('[TR'):
                body.append(lines[j])
            j+=1
        definition=clean(' '.join(body))
        # remove table numeric tails from first year occurrence
        definition=re.split(r'\s(?:\d{2}/\d{2}\s+){2,}', definition)[0].strip()
        if len(definition)>30:
            entries.append({'term':f'{number} {title}','definition':definition,'entry_type':'indicator_definition',
                'source_file':path.name,'source_type':'pdf','dataset_or_file':'Trendrapport HO 2025',
                'field_name':'','values':[],'aliases':[number,title],'tags':infer_tags(title,definition),
                'confidence':'medium'})
        i=max(j,i+1)
    return entries

entries=[]
for p in sorted(BASE.glob('*.docx')):
    entries.extend(extract_docx(p))
entries.extend(parse_txt_1cijfer(BASE/'Bestandsbeschrijving_1cyferho_2025_v1.0.txt'))
entries.extend(parse_txt_vakken(BASE/'Bestandsbeschrijving_Vakkenbestanden.txt'))
entries.extend(parse_handleiding_pdf(BASE/'Handleiding_studentenaantallen_en_studievoortgang_20210809.pdf'))
entries.extend(parse_trendrapport_pdf(BASE/'DUO-trendrapport-ho-2025.pdf'))
# add IDs and consolidate available sources by normalized term later
for idx,e in enumerate(entries,1):
    e['id']=f"{slug(e['term'])}_{idx}"
# Create concept availability map by term normalization
term_map=defaultdict(lambda: {'source_files':set(),'datasets':set(),'field_names':set()})
for e in entries:
    key=slug(re.sub(r'\s*\((?:NIEUW|vanaf 2011|Vanaf 2011).*?\)','',e['term'],flags=re.I))
    term_map[key]['source_files'].add(e.get('source_file',''))
    if e.get('dataset_or_file'): term_map[key]['datasets'].add(e['dataset_or_file'])
    if e.get('field_name'): term_map[key]['field_names'].add(e['field_name'])
for e in entries:
    key=slug(re.sub(r'\s*\((?:NIEUW|vanaf 2011|Vanaf 2011).*?\)','',e['term'],flags=re.I))
    e['available_in_sources']=sorted(term_map[key]['source_files'])
    e['available_in_datasets']=sorted(term_map[key]['datasets'])
    e['related_field_names']=sorted(term_map[key]['field_names'])

# Curated definitions from the extracted docs, manually harmonised to avoid duplicate/conflicting field phrasing
curated=[]
def add_curated(term, definition, source_terms, datasets, aliases=None, fields=None, note=None, tags=None):
    curated.append({
        'id': slug(term), 'term': term, 'definition': clean(definition),
        'aliases': aliases or [], 'related_fields': fields or [], 'available_in_datasets': datasets,
        'source_terms': source_terms, 'note': note or '', 'tags': tags or infer_tags(term, definition)
    })

add_curated('Internationale student',
"Een student wordt als internationale student beschouwd wanneer de student geen Nederlandse nationaliteit heeft en geen Nederlandse vooropleiding voor het hoger onderwijs heeft. In de 1cHO-bestanden wordt dit vastgelegd via indicatievelden. Let op het verschil tussen de actuele variant en de peildatumvariant: bij de actuele variant kan naturalisatie met terugwerkende kracht eerdere jaren wijzigen; bij de peildatumvariant blijft de status voor jaren vóór naturalisatie behouden.",
['Indicatie internationale student','Indicatie internationale student op peildatum 1 oktober','Internationale student (int_student)'],
['1cyferho_2025_v1.0.asc','Inschrijvingen_aggr_UNL_2025.csv','Diplomas_aggr_UNL_2025.csv','EOIcohort_UNL_2025.csv / EOIcohort_21P*_2025.csv','Gediplomeerdencohort_UNL_2025.csv / Gediplomeerdencohort_21P*_2025.csv','VH informatieproducten / 1cijferHO'],
['internationale student','int_student','indicatie internationaal'],
['Indicatie internationale student','Indicatie internationale student op peildatum 1 oktober'],
"De exacte operationalisatie hangt af van het veld: actuele nationaliteit versus nationaliteit op peildatum 1 oktober.")
add_curated('Indicatie internationale student',
"Veld dat aangeeft of een student als internationale student wordt beschouwd. De documentatie vermeldt de waarden J = internationale student en N = geen internationale student. Voor de nationaliteit wordt uitgegaan van de actuele eerste nationaliteit. Daardoor kan een student die later Nederlander wordt in nieuwere 1cHO-bestanden met terugwerkende kracht niet meer als internationale student tellen.",
['Indicatie internationale student'],['1cyferho_2025_v1.0.asc','Inschrijvingen_aggr_UNL_2025.csv','EOIcohort_aggr_UNL_2025.csv','Gediplomeerdencohort_aggr_UNL_2025.csv'],fields=['Indicatie internationale student'])
add_curated('Indicatie internationale student op peildatum 1 oktober',
"Veld dat aangeeft of een student op 1 oktober van het betreffende inschrijvingsjaar als internationale student wordt beschouwd. Hierbij wordt gebruikgemaakt van de eerste nationaliteit op peildatum 1 oktober. Bij latere naturalisatie blijft de student voor jaren vóór naturalisatie als internationale student geregistreerd.",
['Indicatie internationale student op peildatum 1 oktober'],['1cyferho_2025_v1.0.asc','Inschrijvingen_aggr_UNL_2025.csv','Diplomas_aggr_UNL_2025.csv','Gediplomeerdencohort_UNL_2025.csv / Gediplomeerdencohort_21P*_2025.csv'],fields=['Indicatie internationale student op peildatum 1 oktober'])
add_curated('Student / ingeschrevene',
"In de 1cHO-bestanden is een student/ingeschrevene een persoon die met een persoonsgebonden nummer voorkomt in een inschrijvingsrecord. Afhankelijk van de vraag telt niet elke persoon precies één keer: tellingen kunnen per inschrijving, hoofdinschrijving, opleiding, instelling, type hoger onderwijs of cohort worden bepaald.",
['Persoonsgebonden nummer','Inschrijvingen','Inschrijftype'],['1cyferho_2025_v1.0.asc','VH informatieproducten / 1cijferHO'],fields=['Persoonsgebonden nummer','Inschrijvingsvorm','Indicatie actief op peildatum'])
add_curated('Inschrijvingen',
"In overzichten met inschrijvingen wordt elke inschrijving geteld waarbij de student actief was op 1 oktober van het betreffende studiejaar. Standaard worden inschrijvingen ontdubbeld per onderwijstype, waardoor een student met meerdere opleidingen maximaal één keer per onderwijstype wordt meegeteld.",
['Inschrijvingen','Indicatie actief op peildatum','Inschrijftype'],['Inschrijvingen_aggr_UNL_2025.csv','VH informatieproducten / 1cijferHO'],fields=['Indicatie actief op peildatum','Soort inschrijving type ho binnen soort ho','Inschr_type_hbo','Inschr_type_inst','Inschr_type'])
add_curated('Instroom',
"Instroom is een subset van inschrijvingen: een student is actief op 1 oktober en stond sinds 1986 niet eerder op die peildatum ingeschreven binnen het gekozen teldomein. Het teldomein bepaalt of iemand als nieuwe student telt, bijvoorbeeld binnen hoger onderwijs, onderwijssoort, onderwijstype, opleiding of opleiding-instelling.",
['Instroom','Instroomtype','Indicatie eerstejaars'],['Inschrijvingen_aggr_UNL_2025.csv','VH informatieproducten / 1cijferHO','Trendrapport HO 2025'],fields=['Indicatie eerstejaars continu hoger onderwijs','Indicatie eerstejaars continu soort ho','Indicatie eerstejaars continu actuele instelling','Instr_type_hbo','Instr_type_inst','Instr_type'])
add_curated('Diploma’s',
"Diploma’s worden geteld over een volledig studiejaar, niet op een peildatum. In de VH-handleiding worden propedeusediploma’s niet meegeteld. In de 1cHO/UNL-diplomabestanden gaat het om records met een relevant soort diploma/opleidingsfase en een examenresultaat.",
['Diploma’s','Code examenresultaat','Maand examenresultaat'],['Diplomas_aggr_UNL_2025.csv','Gediplomeerdencohort_UNL_2025.csv / Gediplomeerdencohort_21P*_2025.csv','VH informatieproducten / 1cijferHO'],fields=['Code examenresultaat','Maand examenresultaat','Opleidingsfase actueel van het diploma'])
add_curated('EER-student',
"Een EER-student is een student van wie de nationaliteit behoort tot de Europese Economische Ruimte, aangevuld met Zwitserland en Suriname. Nederland is inbegrepen. Bij de peildatumvariant geldt de EER-lijst op 1 oktober van het inschrijvingsjaar; daardoor worden Britse studenten t/m academisch jaar 2021 nog als EER geteld en daarna niet meer.",
['Indicatie EER actueel','Indicatie EER op peildatum 1 oktober'],['1cyferho_2025_v1.0.asc','Inschrijvingen_aggr_UNL_2025.csv','Diplomas_aggr_UNL_2025.csv'],fields=['Indicatie EER actueel','Indicatie EER op peildatum 1 oktober'])
add_curated('Soort hoger onderwijs',
"Geeft aan of de inschrijving valt onder hbo of wo. In VH-informatieproducten is standaard hbo geselecteerd, maar sommige hogescholen bieden ook wo-opleidingen aan.",
['Soort hoger onderwijs of onderwijssoort (soortho)','Soort hoger onderwijs'],['1cyferho_2025_v1.0.asc','VH informatieproducten / 1cijferHO'],fields=['Soort hoger onderwijs'])
add_curated('Type hoger onderwijs binnen soort hoger onderwijs',
"Classificeert de opleiding binnen de onderwijssoort, bijvoorbeeld associate degree, bachelor of master binnen hbo/wo. Dit veld wordt vaak gebruikt als teldomein voor ontdubbeling, instroom en verblijfsjaarberekeningen.",
['Type hoger onderwijs binnen soort hoger onderwijs','Onderwijstype (typeho)'],['1cyferho_2025_v1.0.asc','Inschrijvingen_aggr_UNL_2025.csv','VH informatieproducten / 1cijferHO'],fields=['Type hoger onderwijs binnen soort hoger onderwijs'])
add_curated('Opleidingsvorm',
"Code voor de studievorm waarin de student staat geregistreerd: voltijd, deeltijd of duaal/coöp. In oudere jaren kan het veld soms leeg zijn.",
['Opleidingsvorm'],['1cyferho_2025_v1.0.asc','Inschrijvingen_aggr_UNL_2025.csv','Diplomas_aggr_UNL_2025.csv'],fields=['Opleidingsvorm'])
add_curated('Opleidingsfase',
"Fase van de opleiding waarvoor de student staat ingeschreven of waarvoor het diploma is behaald. De documentatie bevat onder meer codes voor propedeuse, bachelor, master, associate degree, schakelprogramma en oude-stijl fasen.",
['Opleidingsfase','Opleidingsfase actueel','Opleidingsfase actueel van het diploma'],['1cyferho_2025_v1.0.asc','Inschrijvingen_aggr_UNL_2025.csv','Diplomas_aggr_UNL_2025.csv','Gediplomeerdencohort_UNL_2025.csv / Gediplomeerdencohort_21P*_2025.csv'],fields=['Opleidingsfase','Opleidingsfase actueel','Opleidingsfase actueel van het diploma'])
add_curated('Actuele instelling',
"Administratie- of BRIN-code van de instelling zoals die geldig is in het laatst beschikbare inschrijvingsjaar. In de UNL-bestanden worden records beperkt tot 13 door UNL vertegenwoordigde universiteiten.",
['Actuele instelling'],['1cyferho_2025_v1.0.asc','Inschrijvingen_aggr_UNL_2025.csv','Diplomas_aggr_UNL_2025.csv','EOIcohort_UNL_2025.csv / EOIcohort_21P*_2025.csv'],fields=['Actuele instelling'])
add_curated('Opleiding actueel equivalent',
"Actuele opleidingscode/equivalent waarmee historische opleidingen naar een actuele opleiding kunnen worden vertaald. Dit maakt vergelijkingen door de tijd mogelijk.",
['Opleiding actueel equivalent'],['1cyferho_2025_v1.0.asc','Inschrijvingen_aggr_UNL_2025.csv','Diplomas_aggr_UNL_2025.csv'],fields=['Opleiding actueel equivalent'])
add_curated('Opleiding historisch equivalent',
"Historische opleidingsequivalent waarmee records door de tijd heen aan dezelfde historische opleiding kunnen worden gekoppeld. Wordt veel gebruikt in cohortbestanden en ontdubbeling op opleiding-instelling.",
['Opleiding historisch equivalent'],['1cyferho_2025_v1.0.asc','Inschrijvingen_aggr_UNL_2025.csv','EOIcohort_UNL_2025.csv / EOIcohort_21P*_2025.csv'],fields=['Opleiding historisch equivalent'])
add_curated('CROHO-onderdeel actuele opleiding',
"Sector/onderdeel waartoe de actuele opleiding behoort. Mogelijke hoofdwaarden zijn onder andere onderwijs, landbouw/natuurlijke omgeving, natuur, techniek, gezondheidszorg, economie, recht, gedrag en maatschappij, taal en cultuur, en sectoroverstijgend.",
['Croho-onderdeel actuele opleiding'],['1cyferho_2025_v1.0.asc','Inschrijvingen_aggr_UNL_2025.csv','Diplomas_aggr_UNL_2025.csv'],fields=['Croho-onderdeel actuele opleiding'])
add_curated('Verblijfsjaar',
"Aantal jaren dat een student al in een bepaald domein verblijft, bijvoorbeeld hoger onderwijs, soort ho, type ho, actuele opleiding, instelling of opleiding-instelling. Verblijfsjaarvelden zijn belangrijk voor instroom/eerstejaarslogica.",
['Verblijfsjaar hoger onderwijs','Verblijfsjaar soort ho','Verblijfsjaar type ho binnen soort ho','Verblijfsjaar actuele opleiding-instelling'],['1cyferho_2025_v1.0.asc','Inschrijvingen_aggr_UNL_2025.csv'],fields=['Verblijfsjaar hoger onderwijs','Verblijfsjaar soort ho','Verblijfsjaar type ho binnen soort ho','Verblijfsjaar Actuele Opleiding-Instelling'])
add_curated('EOI-cohort',
"EOI staat voor eerste opleiding-instelling/cohort. Het EOI-cohortbestand selecteert studenten vanaf hun eerste inschrijvingsjaar aan een specifieke opleiding-instelling, met filters zoals inschrijvingsjaar vanaf 2011, inschrijvingsvorm S, actief op peildatum en soort inschrijving soort ho 1 t/m 4.",
['EOIcohort_UNL_2025.csv','Eerste jaar aan deze opleiding-instelling'],['EOIcohort_UNL_2025.csv / EOIcohort_21P*_2025.csv','EOIcohort_aggr_UNL_2025.csv'],fields=['Eerste jaar aan deze opleiding-instelling','Voorkomen','Her1 t/m Her8'])
add_curated('Voorkomen',
"In het EOI-cohortbestand geeft Voorkomen het aantal unieke combinaties Actuele instelling + Opleiding historisch equivalent aan dat een student heeft in het EOI-jaar, met correctie voor joint-degree inschrijvingen.",
['Voorkomen'],['EOIcohort_UNL_2025.csv / EOIcohort_21P*_2025.csv','EOIcohort_aggr_UNL_2025.csv'],fields=['Voorkomen'])
add_curated('Her1 t/m Her8',
"Afgeleide herinschrijvings-/statusvelden in EOI-cohortbestanden die met een afleidingsschema bepalen wat de status van de student is in de jaren na het EOI-jaar. Ze gebruiken onder meer persoonsnummer, inschrijvingsjaar, soort ho, instelling, opleiding historisch equivalent, CROHO-onderdeel en opleidingsfase.",
['Her1','Her2','Her3','Her4','Her5','Her6','Her7','Her8'],['EOIcohort_UNL_2025.csv / EOIcohort_21P*_2025.csv','EOIcohort_aggr_UNL_2025.csv'],fields=['Her1','Her2','Her3','Her4','Her5','Her6','Her7','Her8'])
add_curated('Gediplomeerdencohort',
"Cohortbestand gebaseerd op studenten die een relevant diploma hebben behaald. Het bestand wordt gebruikt om diplomering, studieduur en bachelor-masterdoorstroom te analyseren, waaronder directe/indirecte masterinstroom en masterdiploma binnen drie jaar.",
['Gediplomeerdencohort_UNL_2025.csv','Examencohort'],['Gediplomeerdencohort_UNL_2025.csv / Gediplomeerdencohort_21P*_2025.csv','Gediplomeerdencohort_aggr_UNL_2025.csv'],fields=['Masterin','Masterintwee','Masterintot','Masterex3'])
add_curated('Studiesucces',
"Aandeel van een instroomcohort dat na een bepaald aantal jaar een diploma heeft behaald. Standaard kijkt de VH naar studiesucces na vijf en acht jaar voor voltijd bachelorstudenten; voor associate degree/bachelor telt elk einddiploma in het Nederlandse bekostigde hoger onderwijs mee.",
['Studiesucces'],['VH informatieproducten / 1cijferHO','Trendrapport HO 2025'],fields=[])
add_curated('Uitval',
"Aandeel van een instroomcohort dat na een bepaald aantal jaar niet staat ingeschreven in het Nederlands bekostigd hoger onderwijs. Standaard wordt uitval na één jaar en na drie jaar bepaald voor voltijd bachelorstudenten.",
['Uitval'],['VH informatieproducten / 1cijferHO','Trendrapport HO 2025'],fields=[])
add_curated('Studiewissel',
"Aandeel van een instroomcohort dat na een bepaald aantal jaar nog geen diploma heeft behaald, nog wel in het hoger onderwijs staat ingeschreven, maar niet meer bij dezelfde opleiding. Wisselen van instelling met dezelfde studie telt niet als studiewissel.",
['Studiewissel'],['VH informatieproducten / 1cijferHO','Trendrapport HO 2025'],fields=[])
add_curated('Studieduur',
"Gemiddeld aantal maanden dat studenten ingeschreven hebben gestaan in het hbo op het moment van uitstroom, onderscheiden naar uitstroom met of zonder diploma. Alleen maanden waarin de student daadwerkelijk stond ingeschreven worden meegeteld.",
['Studieduur'],['VH informatieproducten / 1cijferHO'],fields=['Maand hoger onderwijs','Maand soort','Maand instelling','Maand equivalent'])
add_curated('Hoogste vooropleiding voor het HO',
"Vooropleiding die relevant is voor de instroom in het hoger onderwijs. Deze wordt gebruikt voor analyses van doorstroom, instroom, internationalisering en studentachtergrond.",
['Hoogste vooropleiding voor het HO','Hoogste vooropl. vóór het HO'],['1cyferho_2025_v1.0.asc','Inschrijvingen_aggr_UNL_2025.csv','Diplomas_aggr_UNL_2025.csv','VH informatieproducten / 1cijferHO'],fields=['Hoogste vooropleiding voor het HO','Vestigingsnummer van de hoogste vooropl. vóór het HO'])
add_curated('Generatie / herkomst',
"Studentachtergrondkenmerken rond migratie/herkomst. In 1cHO 2025 wordt expliciet vermeld dat migratieachtergrond en generatie een verouderde indeling zijn en daarom leeg zijn in het basisbestand; sommige afgeleide/aggregeerde bestanden bevatten nog velden zoals geboorteland en generatie.",
['Generatie','Geboorteland','Herkomstland volgens CBS-definitie'],['1cyferho_2025_v1.0.asc','Inschrijvingen_aggr_UNL_2025.csv','Diplomas_aggr_UNL_2025.csv'],fields=['Geboorteland','Generatie','Herkomstland volgens CBS-definitie','Herkomst-indikking volgens CBS-definitie'])
add_curated('Vakkenbestanden',
"Apart bestand met vakgegevens van havo- en vwo-gediplomeerden van studenten aan de instelling. Het bevat vakcode, omschrijving, gemiddelde cijfers, schoolexamen, centraal examen, eindcijfers, BSN en onderwijsnummer. Het bevat alleen geslaagden havo/vwo aan reguliere vo-instellingen, dus geen vavo en geen particulier vo.",
['Vakkenbestanden'],['Dec_vakcode.asc / vakgegevens'],fields=['Vakcode','Omschrijving vak','Gemiddeld cijfer cijferlijst','Cijfer schoolexamen','Cijfer centraal examen','Burgerservicenummer','Onderwijsnummer'])

# Save files
index_path=OUT/'ho_definities_index.jsonl'
with index_path.open('w',encoding='utf-8') as f:
    for e in entries:
        f.write(json.dumps(e,ensure_ascii=False)+'\n')
curated_path=OUT/'ho_definities_curated.json'
curated_path.write_text(json.dumps({'schema_version':'1.0','purpose':'Automatically cleaned/high-confidence glossary voor conversational AI over 1cHO/HO-documentatie','entries':curated},ensure_ascii=False,indent=2),encoding='utf-8')

# Markdown overview
md=[]
md.append('# HO definities - overzicht voor conversational AI\n')
md.append('Dit bestand is gegenereerd uit de aangeleverde documentatie rond 1cHO/UNL/VH/DUO. Gebruik het als menselijke leeslaag naast de machineleesbare JSON-bestanden. Curated betekent hier automatisch opgeschoonde/high-confidence definities, niet noodzakelijk handmatig goedgekeurde definities.\n')
md.append('## Aanbevolen gebruik in een chatbot\n')
md.append('1. Gebruik `ho_definities_curated.json` voor antwoorden op begripsvragen zoals “wat is een internationale student?”.\n')
md.append('2. Gebruik `ho_definities_index.jsonl` als ruwe RAG-index met veldbeschrijvingen, databestanden en indicator-definities.\n')
md.append('3. Antwoorden over “waar vind ik X?” moeten zoeken in `available_in_datasets`, `related_fields`, `source_file` en `tags`.\n')
md.append('## Curated glossary\n')
for e in curated:
    md.append(f"### {e['term']}\n")
    md.append(e['definition']+'\n')
    if e.get('related_fields'):
        md.append('**Gerelateerde velden:** '+', '.join(e['related_fields'])+'\n')
    if e.get('available_in_datasets'):
        md.append('**Beschikbaar in:** '+', '.join(e['available_in_datasets'])+'\n')
    if e.get('note'):
        md.append('**Let op:** '+e['note']+'\n')
    md.append('')
md.append('## Samenvatting ruwe index\n')
by_type=defaultdict(int)
by_source=defaultdict(int)
for e in entries:
    by_type[e['entry_type']]+=1
    by_source[e['source_file']]+=1
md.append('### Aantallen per type\n')
for k,v in sorted(by_type.items()): md.append(f'- {k}: {v}')
md.append('\n### Aantallen per bron\n')
for k,v in sorted(by_source.items()): md.append(f'- {k}: {v}')
md_path=OUT/'ho_definities_overzicht.md'
md_path.write_text('\n'.join(md),encoding='utf-8')

# Search helper sample
helper = r'''# Voorbeeld: simpele zoekfunctie voor het definitiebestand
import json
from pathlib import Path

CURATED = json.loads(Path("data/ho_definities_curated.json").read_text(encoding="utf-8"))["entries"]
INDEX = [json.loads(line) for line in Path("data/ho_definities_index.jsonl").read_text(encoding="utf-8").splitlines()]

def search_definitions(query, limit=10):
    q = query.lower()
    results = []
    for source_name, rows, weight in [("curated", CURATED, 3), ("index", INDEX, 1)]:
        for row in rows:
            haystack = " ".join(str(row.get(k, "")) for k in ["term", "definition", "aliases", "related_fields", "available_in_datasets", "tags"]).lower()
            score = 0
            for token in q.split():
                if token in haystack:
                    score += weight
                if token in str(row.get("term", "")).lower():
                    score += 2 * weight
            if score:
                results.append((score, source_name, row))
    results.sort(key=lambda x: x[0], reverse=True)
    return results[:limit]

if __name__ == "__main__":
    for score, source, row in search_definitions("waar vind ik internationale studenten"):
        print(f"[{source} | score={score}] {row.get('term')}")
        print(row.get('definition', '')[:400])
        print("Datasets:", row.get('available_in_datasets'))
        print("Fields:", row.get('related_fields') or row.get('related_field_names'))
        print()
'''
helper_path=OUT/'zoek_definities_voorbeeld.py'
helper_path.write_text(helper,encoding='utf-8')

print('entries',len(entries),'curated',len(curated))
print('paths',index_path, curated_path, md_path, helper_path)
# sanity sample international student entries
for e in entries:
    if 'Indicatie internationale student' in e['term'] and e['definition']:
        print('\nSAMPLE', e['source_file'], e['term'])
        print(e['definition'][:600])
        break
